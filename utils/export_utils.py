"""Excel/Word 내보내기 유틸리티."""

import io
from datetime import datetime

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Pt


def export_report_to_excel(report_text: str, context: dict,
                           matrix_data: dict | None = None) -> bytes:
    """보고서와 데이터를 엑셀 파일로 내보내기."""
    wb = Workbook()

    # 시트 1: 보고서 멘트
    ws_report = wb.active
    ws_report.title = "보고서"
    ws_report.column_dimensions["A"].width = 100

    header_font = Font(name="맑은 고딕", bold=True, size=14)
    body_font = Font(name="맑은 고딕", size=11)

    ws_report.cell(row=1, column=1, value=f"{context.get('current_month', '')} 하자보수비 보고서")
    ws_report["A1"].font = header_font

    ws_report.cell(row=2, column=1, value=f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ws_report["A2"].font = Font(name="맑은 고딕", size=9, color="888888")

    for i, line in enumerate(report_text.split("\n"), start=4):
        ws_report.cell(row=i, column=1, value=line).font = body_font

    # 시트 2: 품목×원인 매트릭스 (있는 경우)
    if matrix_data and matrix_data.get("matrix"):
        ws_matrix = wb.create_sheet("품목×원인 매트릭스")
        matrix = matrix_data["matrix"]
        products = matrix_data["products"]
        tags = matrix_data["tags"]

        # 헤더
        header_fill = PatternFill(start_color="1E88E5", end_color="1E88E5", fill_type="solid")
        header_text_font = Font(name="맑은 고딕", bold=True, color="FFFFFF", size=10)

        ws_matrix.cell(row=1, column=1, value="품목군").font = header_text_font
        ws_matrix["A1"].fill = header_fill
        for j, tag in enumerate(tags, start=2):
            cell = ws_matrix.cell(row=1, column=j, value=tag)
            cell.font = header_text_font
            cell.fill = header_fill

        # 데이터
        for i, product in enumerate(products, start=2):
            ws_matrix.cell(row=i, column=1, value=product).font = body_font
            for j, tag in enumerate(tags, start=2):
                val = matrix.get(product, {}).get(tag, 0)
                ws_matrix.cell(row=i, column=j, value=val)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_report_to_word(report_text: str, context: dict) -> bytes:
    """보고서를 Word 파일로 내보내기."""
    doc = Document()

    # 제목
    title = doc.add_heading(
        f"{context.get('current_month', '')} 하자보수비 보고서", level=0
    )
    doc.add_paragraph(
        f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        style="Subtitle",
    )

    # 본문 파싱
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("  -"):
            doc.add_paragraph(line.lstrip("- ").strip(), style="List Bullet")
        elif line.startswith("▶"):
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
